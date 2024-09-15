# import streamlit as st
# import re
# import os
# from io import StringIO
# from docx import Document

# def clean_and_combine_transcript(content, interviewer, interviewee):
#     # Regex pattern to match response number and timestamps like '421\n00:55:08.906 --> 00:55:15.509\n'
#     pattern = re.compile(r'(\d+)\n(\d{2}:\d{2}:\d{2}\.\d{3}) --> (\d{2}:\d{2}:\d{2}\.\d{3})\n')

#     # Split content by matching the pattern
#     segments = re.split(pattern, content)

#     combined_responses = []
#     current_response = []
#     current_speaker = None
#     start_timestamp = None
#     end_timestamp = None
#     first_response_number = None

#     # Iterate through segments, processing timestamps and text
#     for i in range(1, len(segments), 4):
#         response_number = segments[i].strip()  # e.g., '421'
#         start_time = segments[i + 1].strip()   # e.g., '00:55:08.906'
#         end_time = segments[i + 2].strip()     # e.g., '00:55:15.509'
#         line = segments[i + 3].strip()         # e.g., 'Interviewer: So what I asked was...'

#         # Determine the speaker and check if it's a new speaker
#         if line.startswith(interviewer):
#             # If current speaker is interviewee, finalize the interviewee's response
#             if current_speaker == interviewee and current_response:
#                 combined_responses.append(f"{first_response_number}\n{start_timestamp} --> {end_timestamp}\n{interviewee}: " + " ".join(current_response))
#                 current_response = []

#             # Start a new block for the interviewer
#             current_speaker = interviewer
#             if not current_response:  # Set start time and response number only for the first response in the block
#                 first_response_number = response_number
#                 start_timestamp = start_time
#             current_response.append(line[len(interviewer) + 2:].strip())  # Capture the interviewer’s response
#             end_timestamp = end_time  # Continuously update the end time for the last response

#         elif line.startswith(interviewee):
#             # If current speaker is interviewer, finalize the interviewer's response
#             if current_speaker == interviewer and current_response:
#                 combined_responses.append(f"{first_response_number}\n{start_timestamp} --> {end_timestamp}\n{interviewer}: " + " ".join(current_response))
#                 current_response = []

#             # Start a new block for the interviewee
#             current_speaker = interviewee
#             if not current_response:  # Set start time and response number only for the first response in the block
#                 first_response_number = response_number
#                 start_timestamp = start_time
#             current_response.append(line[len(interviewee) + 2:].strip())  # Capture the interviewee’s response
#             end_timestamp = end_time  # Continuously update the end time for the last response

#         else:
#             # Continue appending the ongoing response and update the end timestamp
#             if line.strip():
#                 current_response.append(line.strip())
#                 end_timestamp = end_time  # Update the end time to the latest for ongoing response

#     # Finalize the last response
#     if current_response:
#         combined_responses.append(f"{first_response_number}\n{start_timestamp} --> {end_timestamp}\n{current_speaker}: " + " ".join(current_response))

#     # Return the combined transcript
#     return '\n\n'.join(combined_responses)

# # Function to read a DOCX file
# def read_docx(file):
#     doc = Document(file)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return '\n'.join(full_text)

# st.set_page_config(page_title="Transcript Cleaner", page_icon=":pencil:", layout="centered")
# st.title('Transcript Cleaner')

# uploaded_file = st.file_uploader("Upload your transcript file", type=["vtt", "txt", "docx"])

# if uploaded_file is not None:
#     interviewer = "Interviewer"
#     interviewee = st.text_input("Enter the participant name (eg. P14, P15 etc.)")
    
#     if st.button("**Clean Transcript**", type="primary"):
#         # Determine the file type and read the content accordingly
#         if uploaded_file.type == "text/plain" or uploaded_file.type == "text/vtt":
#             file_content = uploaded_file.read().decode('utf-8')
#         elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             file_content = read_docx(uploaded_file)
#         else:
#             st.error("Unsupported file format.")
#             st.stop()
        
#         # Ensure newline characters are properly retained
#         original_transcript = file_content.replace('\r\n', '\n').replace('\r', '\n')
#         cleaned_transcript = clean_and_combine_transcript(original_transcript, interviewer, interviewee)
        
#         with st.expander("**Cleaned Transcript**"):
#             st.caption("You can edit the cleaned transcript below or select all to copy.")
#             st.text_area("", cleaned_transcript, height=400)
        
#         # Get the original file name and append "_cleaned"
#         original_filename = uploaded_file.name
#         cleaned_filename = os.path.splitext(original_filename)[0] + '_cleaned.txt'
        
#         st.download_button(
#             label="Download cleaned transcript",
#             data=cleaned_transcript,
#             file_name=cleaned_filename,
#             mime='text/plain'
#         )


import streamlit as st
import re
import os
from io import StringIO
from docx import Document

def clean_and_combine_transcript(content, interviewer, interviewee):
    # Regex pattern to match response number and timestamps like '421\n00:55:08.906 --> 00:55:15.509\n'
    pattern = re.compile(r'(\d+)\n(\d{2}:\d{2}:\d{2}\.\d{3}) --> (\d{2}:\d{2}:\d{2}\.\d{3})\n')

    # Split content by matching the pattern
    segments = re.split(pattern, content)

    combined_responses = []
    current_response = []
    current_speaker = None
    start_timestamp = None
    end_timestamp = None
    first_response_number = None

    # Iterate through segments, processing timestamps and text
    for i in range(1, len(segments), 4):
        response_number = segments[i].strip()  # e.g., '421'
        start_time = segments[i + 1].strip()   # e.g., '00:55:08.906'
        end_time = segments[i + 2].strip()     # e.g., '00:55:15.509'
        line = segments[i + 3].strip()         # e.g., 'Interviewer: So what I asked was...'

        # Determine the speaker and check if it's a new speaker
        if line.startswith(interviewer):
            # If current speaker is interviewee, finalize the interviewee's response
            if current_speaker == interviewee and current_response:
                combined_responses.append(f"{first_response_number}\n{start_timestamp} --> {end_timestamp}\n{interviewee}: " + " ".join(current_response))
                current_response = []

            # Start a new block for the interviewer
            current_speaker = interviewer
            if not current_response:  # Set start time and response number only for the first response in the block
                first_response_number = response_number
                start_timestamp = start_time
            current_response.append(line[len(interviewer) + 2:].strip())  # Capture the interviewer’s response
            end_timestamp = end_time  # Continuously update the end time for the last response

        elif line.startswith(interviewee):
            # If current speaker is interviewer, finalize the interviewer's response
            if current_speaker == interviewer and current_response:
                combined_responses.append(f"{first_response_number}\n{start_timestamp} --> {end_timestamp}\n{interviewer}: " + " ".join(current_response))
                current_response = []

            # Start a new block for the interviewee
            current_speaker = interviewee
            if not current_response:  # Set start time and response number only for the first response in the block
                first_response_number = response_number
                start_timestamp = start_time
            current_response.append(line[len(interviewee) + 2:].strip())  # Capture the interviewee’s response
            end_timestamp = end_time  # Continuously update the end time for the last response

        else:
            # Continue appending the ongoing response and update the end timestamp
            if line.strip():
                current_response.append(line.strip())
                end_timestamp = end_time  # Update the end time to the latest for ongoing response

    # Finalize the last response
    if current_response:
        combined_responses.append(f"{first_response_number}\n{start_timestamp} --> {end_timestamp}\n{current_speaker}: " + " ".join(current_response))

    # Return the combined transcript
    return '\n\n'.join(combined_responses)

# Function to read a DOCX file
def read_docx(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

st.set_page_config(page_title="Transcript Cleaner", page_icon=":pencil:", layout="centered")
st.title('Transcript Cleaner')

# Allow the user to choose between uploading a file or pasting the transcript
option = st.radio("Choose input method:", ("Upload a file", "Paste transcript"))

uploaded_file = None
pasted_text = None

if option == "Upload a file":
    uploaded_file = st.file_uploader("Upload your transcript file", type=["vtt", "txt", "docx"])
elif option == "Paste transcript":
    pasted_text = st.text_area("Paste your transcript here")

if uploaded_file is not None or pasted_text:
    interviewer = "Interviewer"
    interviewee = st.text_input("Enter the participant name (eg. P14, P15 etc.)")
    
    if st.button("**Clean Transcript**", type="primary"):
        if uploaded_file is not None:
            # Determine the file type and read the content accordingly
            if uploaded_file.type == "text/plain" or uploaded_file.type == "text/vtt":
                file_content = uploaded_file.read().decode('utf-8')
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                file_content = read_docx(uploaded_file)
            else:
                st.error("Unsupported file format.")
                st.stop()
        else:
            # Use pasted text if no file was uploaded
            file_content = pasted_text
        
        # Ensure newline characters are properly retained
        original_transcript = file_content.replace('\r\n', '\n').replace('\r', '\n')
        cleaned_transcript = clean_and_combine_transcript(original_transcript, interviewer, interviewee)
        
        with st.expander("**Cleaned Transcript**"):
            st.caption("You can edit the cleaned transcript below or select all to copy.")
            st.text_area("", cleaned_transcript, height=400)
        
        # Get the original file name and append "_cleaned" if a file was uploaded, otherwise just name it.
        if uploaded_file:
            original_filename = uploaded_file.name
            cleaned_filename = os.path.splitext(original_filename)[0] + '_cleaned.txt'
        else:
            cleaned_filename = "pasted_transcript_cleaned.txt"
        
        st.download_button(
            label="Download cleaned transcript",
            data=cleaned_transcript,
            file_name=cleaned_filename,
            mime='text/plain'
        )
